import os
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

async def convert_file(inp, ctype):
    name, _ = os.path.splitext(os.path.basename(inp))
    os.makedirs("output", exist_ok=True)
    try:
        if ctype == "pdf_to_txt":
            out = f"output/{name}.txt"
            r = PdfReader(inp)
            with open(out, "w", encoding="utf-8") as f:
                for p in r.pages: f.write(p.extract_text() or "")
            return out
        if ctype == "pdf_to_docx":
            out = f"output/{name}.docx"
            r = PdfReader(inp)
            d = Document()
            for p in r.pages: d.add_paragraph(p.extract_text() or "")
            d.save(out); return out
        if ctype == "docx_to_pdf":
            out = f"output/{name}.pdf"; c = canvas.Canvas(out, pagesize=letter)
            d = Document(inp); y = 750
            for p in d.paragraphs:
                c.drawString(72, y, p.text); y -= 15
                if y < 50: c.showPage(); y = 750
            c.save(); return out
        if ctype == "txt_to_pdf":
            out = f"output/{name}.pdf"; c = canvas.Canvas(out, pagesize=letter)
            y = 750
            with open(inp, "r", encoding="utf-8") as f:
                for line in f:
                    c.drawString(72, y, line.strip()); y -= 15
                    if y < 50: c.showPage(); y = 750
            c.save(); return out
    except Exception as e:
        print("Conversion error:", e); return None
